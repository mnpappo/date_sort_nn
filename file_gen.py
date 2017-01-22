from docx import Document
from faker import Faker 
from docx.enum.text import WD_ALIGN_PARAGRAPH

document = Document()
fake = Faker()

for _ in range(3):
	document.add_heading(fake.sentence(nb_words=6, variable_nb_words=True), level=0)

	p = document.add_heading(fake.name(), level=2)
	p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

	p = document.add_paragraph(fake.address())
	p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

	# p = document.add_paragraph(fake.date_time_this_century(before_now=True, after_now=False, tzinfo=None))
	p = document.add_paragraph(fake.date(pattern="%Y-%m-%d"))
	p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

	document.add_paragraph(fake.paragraphs(nb=4))

	document.add_paragraph(fake.sentence(nb_words=7, variable_nb_words=True), style='ListBullet')
	document.add_paragraph(fake.sentence(nb_words=5, variable_nb_words=True), style='ListBullet')
	document.add_paragraph(fake.sentence(nb_words=9, variable_nb_words=True), style='ListBullet')

	document.add_paragraph(fake.paragraphs(nb=1))
	document.add_heading(fake.sentence(nb_words=9, variable_nb_words=True), level=3)
	document.add_paragraph(fake.paragraphs(nb=5))

	document.add_page_break()



document.save('normal_fake_data.docx')